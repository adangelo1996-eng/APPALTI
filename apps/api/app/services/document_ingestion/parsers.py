from io import BytesIO
from typing import Tuple

from docx import Document as DocxDocument
from pypdf import PdfReader


def extract_text_from_pdf(data: bytes) -> Tuple[str, int]:
  reader = PdfReader(BytesIO(data))
  pages = []
  for page in reader.pages:
      try:
          pages.append(page.extract_text() or "")
      except Exception:
          pages.append("")
  return "\n\n".join(pages), len(reader.pages)


def extract_text_from_docx(data: bytes) -> Tuple[str, int]:
  file_obj = BytesIO(data)
  doc = DocxDocument(file_obj)
  paragraphs = [p.text for p in doc.paragraphs]
  return "\n".join(paragraphs), len(doc.paragraphs)

